"""
Markdown -> DOCX 转换服务

本模块职责
----------
1) 将 Markdown 文本转换为 DOCX（基于 pandoc）。
2) 在未提供自定义模板时，执行统一的三步样式管线：
    - 删除不需要的段落/字符样式
    - 创建并刷新本模块定义的标准样式
    - 逐段应用字体、缩进、表格、图片、代码块等格式
3) 支持 Mermaid 代码块转图片，并在导出 DOCX 前替换为图片引用。

设计说明
--------
- 入口函数：`convert_md_to_docx()`
- 样式定义统一在 ``styles.definitions`` 中维护
- 样式操作委托给 ``styles.docx_adapter``
"""

import traceback
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt, RGBColor

from ..utils import get_logger
from ..utils.markdown_utils import get_md_text
from ..utils.mermaid_utils import replace_mermaid_with_images, cleanup_temp_images, extract_mermaid_blocks
from ..utils.pandoc_utils import pandoc_convert_file
from ..styles.definitions import (
    StyleDefinition,
    STYLE_BY_NAME,
    ALL_STYLES,
    NORMAL,
    TABLE_TEXT as TABLE_CONFIG,
    IMAGE_PARAGRAPH as IMAGE_CONFIG,
    CUSTOM_LIST as CUSTOM_LIST_CONFIG,
    CODE_BLOCK as CODE_CONFIG,
    TOC_1 as TOC1_CONFIG,
    TOC_2 as TOC2_CONFIG,
    TOC_3 as TOC3_CONFIG,
)
from ..styles.docx_adapter import (
    REQUIRED_PARAGRAPH_STYLES,
    REQUIRED_CHARACTER_STYLES,
    get_style_config,
    apply_para_formatting,
    create_or_update_style,
    set_run_fonts,
    _has_num_pr,
    _has_image,
    _is_code_block,
    _is_toc_paragraph,
    _contains_emoji,
)

logger = get_logger(__name__)




# 常量区：图像尺寸换算与页面默认值
# 1pt = 12700 EMU（Office Open XML 单位）
EMU_PER_POINT = 12700.0
DEFAULT_PAGE_WIDTH_PT = 595
DEFAULT_PAGE_HEIGHT_PT = 842
IMAGE_SIDE_MARGIN_PT = 36
IMAGE_TOP_BOTTOM_MARGIN_PT = 72


# ---------------------------------------------------------------------------
# 模块级辅助函数
# ---------------------------------------------------------------------------




def _get_image_limits(doc) -> tuple[float, float]:
    """根据文档页尺寸计算图片可用宽高上限（pt）。"""
    section = doc.sections[0] if doc.sections else None
    page_width = section.page_width.pt if section and section.page_width else DEFAULT_PAGE_WIDTH_PT
    page_height = section.page_height.pt if section and section.page_height else DEFAULT_PAGE_HEIGHT_PT
    max_width = page_width - 2 * IMAGE_SIDE_MARGIN_PT
    max_height = page_height - 2 * IMAGE_TOP_BOTTOM_MARGIN_PT
    return max_width, max_height


def _xml_tag_name(elem) -> str:
    """提取 XML 本地标签名（去掉命名空间前缀）。"""
    return elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag


def _get_extent_size_emu(extent_elem) -> tuple[int, int] | None:
    """读取 `wp:extent` 的宽高（EMU）。读取失败返回 `None`。"""
    cx_val = extent_elem.get(qn("wp:cx")) or extent_elem.get("cx")
    cy_val = extent_elem.get(qn("wp:cy")) or extent_elem.get("cy")
    if not cx_val or not cy_val:
        return None
    return int(cx_val), int(cy_val)


def _set_extent_size_emu(extent_elem, width_emu: int, height_emu: int) -> None:
    """写入 `wp:extent` 的宽高（EMU），兼容命名空间/非命名空间属性。"""
    if extent_elem.get(qn("wp:cx")) is not None:
        extent_elem.set(qn("wp:cx"), str(width_emu))
        extent_elem.set(qn("wp:cy"), str(height_emu))
    else:
        extent_elem.set("cx", str(width_emu))
        extent_elem.set("cy", str(height_emu))


def _scale_extent_if_needed(extent_elem, context_tag: str, max_width_pt: float, max_height_pt: float) -> None:
    """按比例缩放单个图片节点，使其不超过页面可用范围。"""
    size_emu = _get_extent_size_emu(extent_elem)
    if size_emu is None:
        logger.warning("DEBUG: cx_val or cy_val is None/empty")
        return

    current_width_emu, current_height_emu = size_emu
    current_width_pt = current_width_emu / EMU_PER_POINT
    current_height_pt = current_height_emu / EMU_PER_POINT

    logger.info(
        "Image (%s) size: %.0fx%.0fpt, max: %.0fx%.0fpt",
        context_tag,
        current_width_pt,
        current_height_pt,
        max_width_pt,
        max_height_pt,
    )

    scale_factor = 1.0
    if current_width_pt > max_width_pt:
        scale_factor = min(scale_factor, max_width_pt / current_width_pt)
    if current_height_pt > max_height_pt:
        scale_factor = min(scale_factor, max_height_pt / current_height_pt)

    if scale_factor >= 1.0:
        logger.info("Image within limits, no scaling needed")
        return

    new_width_emu = int(current_width_emu * scale_factor)
    new_height_emu = int(current_height_emu * scale_factor)
    _set_extent_size_emu(extent_elem, new_width_emu, new_height_emu)

    logger.info(
        "✓ Scaled image from %.0fx%.0fpt to %.0fx%.0fpt",
        current_width_pt,
        current_height_pt,
        new_width_emu / EMU_PER_POINT,
        new_height_emu / EMU_PER_POINT,
    )


def _scale_images_in_paragraph(paragraph, max_width_pt: float, max_height_pt: float) -> None:
    """扫描并缩放段落内全部图片节点。"""
    logger.info("Found image paragraph, checking scaling...")
    current_context_tag = "unknown"

    try:
        for elem in paragraph._element.iter():
            tag_name = _xml_tag_name(elem)

            if tag_name in {"inline", "anchor"}:
                current_context_tag = tag_name
                continue

            if tag_name != "extent":
                continue

            try:
                _scale_extent_if_needed(elem, current_context_tag, max_width_pt, max_height_pt)
            except Exception as img_exc:
                logger.warning(f"Failed to scale image: {img_exc}")
                logger.warning(f"Traceback: {traceback.format_exc()}")
    except Exception as exc:
        logger.error(f"Error iterating paragraph elements: {exc}")
        logger.error(f"Traceback: {traceback.format_exc()}")


def _apply_table_layout(tbl) -> None:
    """为表格统一设置宽度、布局模式与边框样式。

    - 宽度：100% 页面宽（`w:type=pct, w=5000`）
    - 布局：autofit，列宽根据内容自动调整
    - 边框：外边框与内部网格均为 single
    """
    tblPr = tbl._element.tblPr
    if tblPr is None:
        tblPr = parse_xml(f"<w:tblPr {nsdecls('w')}/>")
        tbl._element.insert(0, tblPr)

    existing_tblW = tblPr.find(qn("w:tblW"))
    if existing_tblW is not None:
        tblPr.remove(existing_tblW)

    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>')
    tblPr.append(tblW)

    # 设置表格布局为 autofit，使列宽根据内容自动调整
    existing_tblLayout = tblPr.find(qn("w:tblLayout"))
    if existing_tblLayout is not None:
        tblPr.remove(existing_tblLayout)

    tblLayout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:val="autofit"/>')
    tblPr.append(tblLayout)

    tblBorders_xml = f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>
    </w:tblBorders>'''
    tblBorders = parse_xml(tblBorders_xml)

    existing_tblBorders = tblPr.find(qn("w:tblBorders"))
    if existing_tblBorders is not None:
        tblPr.remove(existing_tblBorders)

    tblPr.append(tblBorders)


def _set_cell_vertical_center(cell) -> None:
    """设置单元格垂直居中。"""
    tc = cell._element
    tcPr = tc.tcPr if tc.tcPr is not None else parse_xml(f'<w:tcPr {nsdecls("w")}/>')
    if tc.tcPr is None:
        tc.insert(0, tcPr)

    existing_vAlign = tcPr.find(qn("w:vAlign"))
    if existing_vAlign is not None:
        tcPr.remove(existing_vAlign)

    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)


# 表格前后间距常量（半行高度），单位：缇（twips），1 pt = 20 twips
_HALF_LINE_TWIPS = 120  # 6 pt，约半行高度


def _adjust_paragraph_spacing(p_elem, space_before_twips=None, space_after_twips=None) -> None:
    """在段落现有间距基础上累加指定值。

    若段落尚无 spacing 元素则新建；否则在现有值上累加。
    用于在表格前后增加半行间距，不丢失已有间距信息。

    Args:
        p_elem: ``w:p`` XML 元素
        space_before_twips: 要累加的段前间距（缇）
        space_after_twips: 要累加的段后间距（缇）
    """
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_elem.insert(0, pPr)

    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)

    if space_before_twips is not None:
        existing = spacing.get(qn("w:before"))
        new_val = (int(existing) + space_before_twips) if existing else space_before_twips
        spacing.set(qn("w:before"), str(new_val))

    if space_after_twips is not None:
        existing = spacing.get(qn("w:after"))
        new_val = (int(existing) + space_after_twips) if existing else space_after_twips
        spacing.set(qn("w:after"), str(new_val))


def _add_spacing_around_table(tbl) -> None:
    """在表格前后各增加半行间距，避免表格与正文紧贴。

    利用 XML 兄弟节点关系找到表格前后的段落（``w:p``），
    为前一段增加段后间距，为后一段增加段前间距。

    Args:
        tbl: ``python-docx`` Table 对象
    """
    tbl_elem = tbl._element

    # 前一个兄弟段落 → 增加段后间距
    prev = tbl_elem.getprevious()
    if prev is not None and prev.tag == qn("w:p"):
        _adjust_paragraph_spacing(prev, space_after_twips=_HALF_LINE_TWIPS)

    # 后一个兄弟段落 → 增加段前间距
    next_elem = tbl_elem.getnext()
    if next_elem is not None and next_elem.tag == qn("w:p"):
        _adjust_paragraph_spacing(next_elem, space_before_twips=_HALF_LINE_TWIPS)


def _format_table_content(doc) -> None:
    """对文档中所有表格应用统一格式。

    包含：
    - 表格级（宽度、边框、前后间距）
    - 单元格级（垂直居中）
    - 段落级（`Table Text` 样式 + run 字体设置）
    """
    table_text_style = doc.styles["Table Text"]
    for tbl in doc.tables:
        _apply_table_layout(tbl)
        _add_spacing_around_table(tbl)
        for row in tbl.rows:
            for cell in row.cells:
                _set_cell_vertical_center(cell)
                for para in cell.paragraphs:
                    try:
                        para.style = table_text_style
                        apply_para_formatting(para, TABLE_CONFIG, is_table=True)
                    except Exception as exc:
                        logger.warning(f"Failed to format table cell: {exc}")




# ---------------------------------------------------------------------------
# 三步格式化管线
# ---------------------------------------------------------------------------


def _step1_delete_styles(doc) -> None:
    """步骤 1：删除不在白名单中的样式，并做必要迁移。

    关键点：
    - 段落样式：不在 `REQUIRED_PARAGRAPH_STYLES` 中的将被移除；
    - 字符样式：不在 `REQUIRED_CHARACTER_STYLES` 中的将被移除；
    - 删除前先把引用它们的段落/run 重映射，避免悬空引用。
    """
    normal_style = doc.styles["Normal"]
    existing_names = {s.name for s in doc.styles}
    # List Paragraph may not exist yet before _step2 creates it; create a minimal
    # placeholder now so that list sub-styles can be reassigned to it.
    if "List Paragraph" not in existing_names:
        list_style = doc.styles.add_style("List Paragraph", 1)
        list_style.base_style = normal_style
    else:
        list_style = doc.styles["List Paragraph"]

    styles_to_delete = [
        s
        for s in doc.styles
        if (s.type == 1 and s.name not in REQUIRED_PARAGRAPH_STYLES)
        or (s.type == 2 and s.name not in REQUIRED_CHARACTER_STYLES)
    ]

    for style_obj in styles_to_delete:
        try:
            if style_obj.type == 1:
                # Pandoc generates "List Paragraph 1/2/..." for nested lists;
                # reassign those to List Paragraph so they keep bullet formatting.
                is_list_sub = "List" in style_obj.name
                fallback_style = list_style if is_list_sub else normal_style
                for para in doc.paragraphs:
                    if para.style.name == style_obj.name:
                        para.style = fallback_style
                for tbl in doc.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if para.style.name == style_obj.name:
                                    para.style = fallback_style
            else:
                style_id = style_obj._element.get(qn("w:styleId"))
                all_paras = list(doc.paragraphs)
                for tbl in doc.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            all_paras.extend(cell.paragraphs)
                for para in all_paras:
                    for run in para.runs:
                        rPr = run._element.find(qn("w:rPr"))
                        if rPr is not None:
                            rStyle = rPr.find(qn("w:rStyle"))
                            if rStyle is not None and rStyle.get(qn("w:val")) == style_id:
                                rPr.remove(rStyle)

            style_obj._element.getparent().remove(style_obj._element)
            logger.info(f"Deleted style: {style_obj.name}")
        except Exception as exc:
            logger.warning(f"Failed to delete style '{style_obj.name}': {exc}")


def _step2_create_styles(doc) -> None:
    """步骤 2：创建/刷新本模块定义的标准样式。"""
    for style_def in ALL_STYLES:
        try:
            create_or_update_style(doc, style_def)
            logger.debug(f"Created/updated style: {style_def.name}")
        except Exception as exc:
            logger.warning(f"Failed to create/update style '{style_def.name}': {exc}")


def _step3_apply_to_content(doc) -> None:
    """步骤 3：遍历正文并按内容类型应用样式。"""
    image_para_style = doc.styles["Image Paragraph"]
    custom_list_style = doc.styles["Custom List"]
    code_block_style = doc.styles["Code Block"]
    toc1_style = doc.styles["Table of Contents 1"]
    toc2_style = doc.styles["Table of Contents 2"]
    toc3_style = doc.styles["Table of Contents 3"]

    code_block_count = 0
    toc_count = {1: 0, 2: 0, 3: 0}
    max_image_width, max_image_height = _get_image_limits(doc)

    for para in doc.paragraphs:
        try:
            if _is_code_block(para):
                para.style = code_block_style
                apply_para_formatting(para, CODE_CONFIG)
                code_block_count += 1
                logger.debug(f"Applied Code Block style to paragraph: {para.text[:50]}")
            elif _has_image(para):
                para.style = image_para_style
                apply_para_formatting(para, IMAGE_CONFIG)
                _scale_images_in_paragraph(para, max_image_width, max_image_height)
            elif _has_num_pr(para):
                para.style = custom_list_style
                apply_para_formatting(para, CUSTOM_LIST_CONFIG)
            else:
                # 检查是否为目录项
                is_toc, toc_level = _is_toc_paragraph(para)
                if is_toc:
                    toc_style_map = {1: toc1_style, 2: toc2_style, 3: toc3_style}
                    toc_config_map = {1: TOC1_CONFIG, 2: TOC2_CONFIG, 3: TOC3_CONFIG}
                    para.style = toc_style_map[toc_level]
                    apply_para_formatting(para, toc_config_map[toc_level])
                    toc_count[toc_level] += 1
                else:
                    style_name = para.style.name if para.style else "Normal"
                    style_def = get_style_config(style_name)
                    apply_para_formatting(para, style_def)
        except Exception as exc:
            logger.warning(f"Failed to format paragraph: {exc}")

    logger.info(f"Total code blocks formatted: {code_block_count}")
    logger.info(f"Total TOC items formatted: Level 1: {toc_count[1]}, Level 2: {toc_count[2]}, Level 3: {toc_count[3]}")
    _format_table_content(doc)


def _build_pandoc_extra_args(
    is_enable_toc: bool,
    resource_paths: list[str] | None = None,
) -> list[str]:
    """构建 Pandoc 参数列表。

    包含：
    - `--toc`：目录开关
    - `--resource-path`：资源搜索路径（图片等）
    """
    extra_args: list[str] = []
    if is_enable_toc:
        extra_args.append("--toc")
    if resource_paths:
        extra_args.append(f"--resource-path={';'.join(resource_paths)}")
    return extra_args


def _convert_markdown_file_to_docx(
    source_md_file: Path,
    output_path: Path,
    is_enable_toc: bool,
    resource_paths: list[str] | None = None,
) -> None:
    """将单个 Markdown 文件转换为 DOCX，并应用默认格式化。"""
    extra_args = _build_pandoc_extra_args(is_enable_toc, resource_paths)

    pandoc_convert_file(
        source_file=str(source_md_file),
        input_format="markdown",
        dest_format="docx",
        outputfile=str(output_path),
        extra_args=extra_args,
    )

    _apply_formatting(output_path)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def _apply_formatting(docx_path: Path) -> None:
    """
        对输出 DOCX 应用三步格式化管线：
        1) 删除非白名单样式；
        2) 依据 `ALL_STYLES` 刷新标准样式；
        3) 遍历正文/表格并应用段落与 run 级格式。
    """
    doc = Document(docx_path)
    _step1_delete_styles(doc)
    _step2_create_styles(doc)
    _step3_apply_to_content(doc)
    doc.save(docx_path)
    logger.info(f"Applied formatting to {docx_path}")


def convert_md_to_docx(
    md_text: str,
    output_path: Path,
    is_strip_wrapper: bool = False,
    is_enable_toc: bool = False,
    convert_mermaid: bool = True,
    save_mermaid_images: bool = False,
    output_dir: Path | None = None,
) -> None:
    """
    将 Markdown 文本转换为 DOCX。

    Args:
        md_text: 输入 Markdown 文本
        output_path: 输出 DOCX 路径
        is_strip_wrapper: 是否剥离代码块包装
        is_enable_toc: 是否启用目录
        convert_mermaid: 是否启用 Mermaid 转图片
        save_mermaid_images: 是否将 Mermaid 图片持久保存到输出目录
        output_dir: Mermaid 图片输出目录（当 `save_mermaid_images=True` 时建议提供）

    Raises:
        ValueError: 输入处理失败
        Exception: 转换或后处理失败
    """
    processed_md = get_md_text(md_text, is_strip_wrapper=is_strip_wrapper)

    # 预扫描 Mermaid 代码块，用于分流到“图表转换流程”或“标准流程”
    mermaid_blocks = extract_mermaid_blocks(processed_md)

    if mermaid_blocks and convert_mermaid:
        logger.info(f"检测到 {len(mermaid_blocks)} 个 Mermaid 图表，开始转换...")

        # 根据保存策略决定图片落盘位置
        if save_mermaid_images and output_dir:
            # 创建图片保存目录
            images_dir = output_dir / "mermaid_images"
            images_dir.mkdir(exist_ok=True)
            save_path = images_dir
            logger.info(f"Mermaid 图片将保存到: {images_dir}")
        else:
            # 不保存到目标目录：统一写临时目录，流程结束后清理
            save_path = None

        # 统一使用临时目录承载中间产物（Markdown、转换图片等）
        with TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)

            # 若不持久化 Mermaid 图片，则直接写入临时目录
            image_save_path = save_path if save_mermaid_images else temp_path

            # 替换 Mermaid 代码块为图片引用（使用 PNG 格式，通过 scale 参数提高清晰度）
            modified_md, generated_images, mermaid_stats, failed_mermaid_codes = replace_mermaid_with_images(
                processed_md,
                image_save_path,
                image_format="png",
                timeout=15,  # 增加超时时间，因为大图片需要更长时间
                max_retries=3,
                retry_delay=2,
                scale=3  # 3倍缩放提高清晰度
            )
            
            # 如果有失败的 Mermaid 代码，保存到临时文件
            if failed_mermaid_codes:
                failed_codes_file = output_path.parent / f"{output_path.stem}_failed_mermaid_codes.md"
                with open(failed_codes_file, 'w', encoding='utf-8') as f:
                    f.write(f"# 转换失败的 Mermaid 代码\n\n")
                    f.write(f"共计 {len(failed_mermaid_codes)} 个 Mermaid 图表转换失败\n\n")
                    f.write("=" * 60 + "\n\n")
                                
                    for item in failed_mermaid_codes:
                        f.write(f"## Mermaid 图表 #{item['index']}\n\n")
                        f.write(f"**预期文件名**: {item['filename']}\n\n")
                        f.write(f"**代码**:\n\n```mermaid\n{item['code']}\n```\n\n")
                        f.write("-" * 60 + "\n\n")
                            
                logger.info(f"已将 {len(failed_mermaid_codes)} 个失败的 Mermaid 代码保存到: {failed_codes_file.name}")

            # 将替换后的 Markdown（Mermaid -> ![](...)）写到临时文件
            temp_md_file = temp_path / "temp.md"
            temp_md_file.write_text(modified_md, encoding="utf-8")

            # 指定 Pandoc 资源路径，确保图片可被解析
            resource_paths = [str(temp_path)]
            if save_mermaid_images and save_path:
                resource_paths.append(str(save_path))

            try:
                _convert_markdown_file_to_docx(
                    source_md_file=temp_md_file,
                    output_path=output_path,
                    is_enable_toc=is_enable_toc,
                    resource_paths=resource_paths,
                )
            finally:
                # 非持久化模式：主动清理 Mermaid 图片临时文件
                if not save_mermaid_images:
                    cleanup_temp_images(generated_images)
                    logger.info("已清理临时图片文件")
                else:
                    logger.info(f"已保存 {len(generated_images)} 个 Mermaid 图片到: {save_path}")
        
        # 在文档转换完成后输出 Mermaid 汇总
        if mermaid_stats and mermaid_stats["total"] > 0:
            logger.info("=" * 50)
            logger.info("Mermaid 转换汇总:")
            logger.info(f"  总计: {mermaid_stats['total']} 个")
            logger.info(f"  成功: {mermaid_stats['success']} 个")
            if mermaid_stats["failed"] > 0:
                logger.info(f"  失败: {mermaid_stats['failed']} 个")
            logger.info("=" * 50)
    else:
        # 无 Mermaid：走标准 Markdown -> DOCX 流程
        logger.info("未检测到 Mermaid 图表，使用标准转换流程")

        with TemporaryDirectory() as temp_dir:
            temp_md_file = Path(temp_dir) / "temp.md"
            temp_md_file.write_text(processed_md, encoding="utf-8")
            _convert_markdown_file_to_docx(
                source_md_file=temp_md_file,
                output_path=output_path,
                is_enable_toc=is_enable_toc,
            )



